import argparse
import json
from pathlib import Path

from docx import Document


def parse_args():
    parser = argparse.ArgumentParser(description="Initialize an auto-lab run directory.")
    parser.add_argument("--requirements", required=True, help="Path to the requirement document.")
    parser.add_argument("--template", required=True, help="Path to the docx template.")
    parser.add_argument("--output-dir", required=True, help="Directory for the generated run files.")
    parser.add_argument("--output-docx-name", required=True, help="Final output docx file name.")
    return parser.parse_args()


def analyze_template(template_path: Path):
    doc = Document(str(template_path))
    paragraphs = list(doc.paragraphs)
    non_empty = [p.text.strip() for p in paragraphs if p.text.strip()]
    table_shapes = []
    for table in doc.tables:
        rows = len(table.rows)
        cols = max((len(row.cells) for row in table.rows), default=0)
        table_shapes.append({"rows": rows, "cols": cols})

    return {
        "paragraph_count": len(paragraphs),
        "non_empty_paragraph_count": len(non_empty),
        "table_count": len(doc.tables),
        "table_shapes": table_shapes,
        "sample_headings_or_text": non_empty[:20],
    }


def write_script_stub(path: Path, script_role: str):
    stub = f'''import argparse
from pathlib import Path

AUTO_LAB_TEMPLATE_SCRIPT_STUB = True


def parse_args():
    parser = argparse.ArgumentParser(description="Template-specific {script_role} script for the current docx task.")
'''
    if script_role == "fill":
        stub += """    parser.add_argument("--template", required=True)
    parser.add_argument("--copywriting", required=True)
    parser.add_argument("--output", required=True)
"""
    elif script_role == "insert":
        stub += """    parser.add_argument("--docx", required=True)
    parser.add_argument("--insert-config", required=True)
"""
    else:
        stub += """    parser.add_argument("--docx", required=True)
    parser.add_argument("--insert-config", required=True)
"""
    stub += f"""    return parser.parse_args()


def main():
    args = parse_args()
    raise SystemExit(
        "This is a template-specific {script_role} script stub. "
        "Analyze the current template and replace this stub with logic that preserves the original document structure."
    )


if __name__ == "__main__":
    main()
"""
    path.write_text(stub, encoding="utf-8")


def main():
    args = parse_args()

    requirements = Path(args.requirements).expanduser().resolve()
    template = Path(args.template).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    images_dir = output_dir / "generated_images"
    images_dir.mkdir(exist_ok=True)
    task_scripts_dir = output_dir / "task_scripts"
    task_scripts_dir.mkdir(exist_ok=True)

    output_docx = output_dir / args.output_docx_name
    workflow_path = output_dir / "workflow.json"
    copywriting_path = output_dir / "配文.md"
    prompt_config_path = output_dir / "prompt_config.json"
    insert_config_path = output_dir / "insert_config.json"
    template_manifest_path = output_dir / "template_manifest.json"
    fill_script_path = task_scripts_dir / "fill_template.py"
    insert_script_path = task_scripts_dir / "insert_images.py"
    verify_script_path = task_scripts_dir / "verify_template.py"

    template_manifest = analyze_template(template)
    template_manifest["template_path"] = str(template)
    template_manifest["template_protection_rule"] = "Do not change original structure, fixed text, or template styling unless a position is explicitly identified as fillable."

    workflow = {
        "requirements_path": str(requirements),
        "template_path": str(template),
        "output_dir": str(output_dir),
        "output_docx": str(output_docx),
        "copywriting_path": str(copywriting_path),
        "prompt_config_path": str(prompt_config_path),
        "insert_config_path": str(insert_config_path),
        "images_dir": str(images_dir),
        "template_manifest_path": str(template_manifest_path),
        "task_scripts_dir": str(task_scripts_dir),
        "docx_scripts": {
            "fill": str(fill_script_path),
            "insert": str(insert_script_path),
            "verify": str(verify_script_path)
        },
        "commands": {
            "fill_command": 'python "{fill_script}" --template "{template}" --copywriting "{copywriting}" --output "{output_docx}"',
            "insert_command": 'python "{insert_script}" --docx "{output_docx}" --insert-config "{insert_config}"',
            "verify_command": 'python "{verify_script}" --docx "{output_docx}" --insert-config "{insert_config}"'
        }
    }

    prompt_config = {
        "total_count": 0,
        "resolution": "2560x1440",
        "output_dir": str(images_dir),
        "max_workers": 3,
        "max_retries": 3,
        "retry_delay": 2,
        "image_policy": {
            "default_mode": "screenshot_strict",
            "auto_append_negative": True,
            "fail_on_prompt_risk": True,
            "forbidden_terms": [
                "流程图",
                "架构图",
                "讲解框",
                "说明面板",
                "悬浮标注",
                "箭头标注",
                "海报",
                "poster",
                "callout",
                "annotation",
                "flowchart",
                "diagram"
            ]
        },
        "global_prompt": "",
        "images": []
    }

    insert_config = {
        "target_docx": str(output_docx),
        "images": {}
    }

    copywriting = """# 配文\n\n> 在这里先写正文，再插入图片占位符。\n>\n> 占位符格式示例：`{{img_xx}}`\n>\n> 占位符必须独占一行，并且同一个 key 需要同时出现在 prompt_config.json 和 insert_config.json 中。\n"""

    write_script_stub(fill_script_path, "fill")
    write_script_stub(insert_script_path, "insert")
    write_script_stub(verify_script_path, "verify")

    workflow_path.write_text(json.dumps(workflow, ensure_ascii=False, indent=2), encoding="utf-8")
    prompt_config_path.write_text(json.dumps(prompt_config, ensure_ascii=False, indent=2), encoding="utf-8")
    insert_config_path.write_text(json.dumps(insert_config, ensure_ascii=False, indent=2), encoding="utf-8")
    copywriting_path.write_text(copywriting, encoding="utf-8")
    template_manifest_path.write_text(json.dumps(template_manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"Initialized run directory: {output_dir}")
    print(f"  workflow: {workflow_path}")
    print(f"  copywriting: {copywriting_path}")
    print(f"  prompt config: {prompt_config_path}")
    print(f"  insert config: {insert_config_path}")
    print(f"  template manifest: {template_manifest_path}")
    print(f"  task scripts: {task_scripts_dir}")


if __name__ == "__main__":
    main()
